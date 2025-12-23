#!/usr/bin/env python3
"""
Script to create a well-formatted DOCX from the DZONE_ARTICLE.md
with proper diagram formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import os

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_architecture_diagram(doc):
    """Create a properly formatted architecture diagram as a table."""

    # Add a heading for the diagram
    doc.add_paragraph()

    # Create outer table (Browser container)
    outer_table = doc.add_table(rows=2, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Browser header
    header_cell = outer_table.rows[0].cells[0]
    header_cell.text = "Browser"
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    set_cell_shading(header_cell, "E8E8E8")

    # Content cell with nested table
    content_cell = outer_table.rows[1].cells[0]

    # Create nested table for Sidebar and Redoc Viewer
    nested_table = content_cell.add_table(rows=1, cols=3)
    nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Sidebar cell
    sidebar_cell = nested_table.rows[0].cells[0]
    sidebar_content = """Sidebar

Category A
  • API 1
  • API 2
Category B
  • API 3"""
    sidebar_cell.text = sidebar_content
    set_cell_shading(sidebar_cell, "F5F5F5")
    for para in sidebar_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Arrow cell
    arrow_cell = nested_table.rows[0].cells[1]
    arrow_cell.text = "→"
    arrow_para = arrow_cell.paragraphs[0]
    arrow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = arrow_para.runs[0]
    run.font.size = Pt(18)
    run.bold = True

    # Redoc Viewer cell
    redoc_cell = nested_table.rows[0].cells[2]
    redoc_content = """Redoc Viewer

• API Info
• Endpoints
• Request/Response Schemas
• Code Samples"""
    redoc_cell.text = redoc_content
    set_cell_shading(redoc_cell, "F5F5F5")
    for para in redoc_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths
    nested_table.columns[0].width = Inches(2.0)
    nested_table.columns[1].width = Inches(0.5)
    nested_table.columns[2].width = Inches(3.0)

    # Add borders to outer table
    for row in outer_table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    for row in nested_table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    doc.add_paragraph()
    return outer_table

def set_cell_border(cell, border_color="000000", border_size="4"):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_tech_stack_table(doc):
    """Create a well-formatted technology stack table."""
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Technology", "Purpose"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        set_cell_shading(cell, "3B82F6")
        run.font.color.rgb = RGBColor(255, 255, 255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    data = [
        ("React 18", "UI framework"),
        ("TypeScript", "Type safety"),
        ("Vite", "Lightning-fast build tool"),
        ("Redoc", "API documentation rendering"),
        ("Tailwind CSS", "Styling"),
        ("webapi-parser", "RAML-to-OpenAPI conversion"),
    ]

    for i, (tech, purpose) in enumerate(data, start=1):
        table.rows[i].cells[0].text = tech
        table.rows[i].cells[1].text = purpose

        # Alternate row colors
        if i % 2 == 0:
            set_cell_shading(table.rows[i].cells[0], "F3F4F6")
            set_cell_shading(table.rows[i].cells[1], "F3F4F6")

    # Set column widths
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(3.5)

    # Add borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    doc.add_paragraph()
    return table

def add_references_section(doc):
    """Add the references section to the document."""

    # Page break before references
    doc.add_page_break()

    # References heading
    doc.add_heading('References', level=1)

    # Official Documentation section
    doc.add_heading('Official Documentation & Specifications', level=2)

    references_official = [
        ("OpenAPI Specification v3.2.0", "https://spec.openapis.org/oas/v3.2.0.html"),
        ("OpenAPI Initiative", "https://www.openapis.org/"),
        ("OpenAPI Specification GitHub Repository", "https://github.com/OAI/OpenAPI-Specification"),
        ("RAML Official Website", "https://raml.org/"),
        ("RAML 1.0 Specification", "https://github.com/raml-org/raml-spec/blob/master/versions/raml-10/raml-10.md"),
        ("Redoc GitHub Repository", "https://github.com/Redocly/redoc"),
        ("Redoc Official Documentation", "https://redocly.com/docs/redoc"),
        ("React Official Documentation", "https://react.dev/"),
        ("Vite Official Documentation", "https://vite.dev/"),
        ("Tailwind CSS Documentation", "https://tailwindcss.com/docs"),
        ("GitHub Actions Documentation", "https://docs.github.com/en/actions"),
        ("webapi-parser (npm)", "https://www.npmjs.com/package/webapi-parser"),
    ]

    for i, (title, url) in enumerate(references_official, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(f"{title}")
        run.bold = True
        para.add_run(f"\n{url}")

    # Academic Research section
    doc.add_heading('Academic & Research Papers', level=2)

    references_academic = [
        ("Meng, M., Steinhardt, S., & Schubert, A. (2019)",
         "How Developers Use API Documentation: An Observation Study",
         "Communication Design Quarterly, Vol 7, No. 2. ACM SIGDOC",
         "https://dl.acm.org/doi/10.1145/3358931.3358937"),
        ("Henkel, M., & Keshishzadeh, S. (2020)",
         "Optimizing API Documentation",
         "ACM SIGDOC Annual International Conference on Design of Communication",
         "https://dl.acm.org/doi/fullHtml/10.1145/3380851.3416759"),
        ("IEEE/ACM ICSE (2024)",
         "Managing API Evolution in Microservice Architecture",
         "46th International Conference on Software Engineering",
         "https://dl.acm.org/doi/10.1145/3639478.3639800"),
        ("Journal of Systems and Software (2024)",
         "Microservice API Evolution in Practice: A Study on Strategies and Challenges",
         "Elsevier",
         "https://www.sciencedirect.com/science/article/pii/S0164121224001559"),
        ("Di Francesco, P., Malavolta, I., & Lago, P. (2021)",
         "On Microservice Analysis and Architecture Evolution: A Systematic Mapping Study",
         "MDPI Applied Sciences",
         "https://www.mdpi.com/2076-3417/11/17/7856"),
        ("Heinrich, R., et al. (2017)",
         "Performance Engineering for Microservices: Research Challenges and Directions",
         "ACM/SPEC International Conference on Performance Engineering",
         "https://dl.acm.org/doi/10.1145/3053600.3053653"),
    ]

    for i, (authors, title, venue, url) in enumerate(references_academic, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(f"{authors}. ")
        para.add_run(f'"{title}." ')
        run2 = para.add_run(f"{venue}.")
        run2.italic = True
        para.add_run(f"\n{url}")

    # Industry Reports section
    doc.add_heading('Industry Reports & Best Practices', level=2)

    references_industry = [
        ("SmartBear", "State of Software Quality - API Report 2023",
         "https://smartbear.com/state-of-software-quality/api/"),
        ("Postman", "API Documentation: How to Write, Examples & Best Practices",
         "https://www.postman.com/api-platform/api-documentation/"),
        ("Swagger/SmartBear", "API Documentation: The Secret to a Great API Developer Experience",
         "https://swagger.io/resources/ebooks/api-documentation-the-secret-to-a-great-api-developer-experience/"),
        ("Pronovix", "Developer Experience Best Practices - API The Docs 2023",
         "https://pronovix.com/articles/developer-experience-best-practices-api-docs-2023"),
    ]

    for i, (org, title, url) in enumerate(references_industry, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(f"{org}. ")
        run.bold = True
        para.add_run(f'"{title}."')
        para.add_run(f"\n{url}")

    # Key Statistics section
    doc.add_heading('Key Statistics', level=2)

    stats = [
        "64% of developers express frustration when faced with poor API documentation resources",
        "Documentation with interactive components reduces support queries by up to 30%",
        "Developers experienced a 40% decrease in onboarding time following documentation improvements",
        "Support requests fell by an average of 35% with improved documentation",
        "Satisfaction scores improved from 60% to over 85% after documentation enhancements",
    ]

    for stat in stats:
        para = doc.add_paragraph(stat, style='List Bullet')


def process_markdown_to_docx(md_path, docx_path):
    """Convert markdown to docx with proper formatting."""

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    code_lang = ""
    skip_until_code_end = False

    while i < len(lines):
        line = lines[i]

        # Check for the ASCII diagram section and replace it
        if '┌─────────────────────────────────────────────────────┐' in line:
            # Skip the ASCII diagram
            while i < len(lines) and '└─────────────────────────────────────────────────────┘' not in lines[i]:
                i += 1
            i += 1  # Skip the closing line
            if i < len(lines) and lines[i] == '```':
                i += 1  # Skip the closing code fence

            # Insert the formatted diagram
            create_architecture_diagram(doc)
            continue

        # Check for technology stack table
        if '| Technology | Purpose |' in line:
            # Skip the markdown table
            while i < len(lines) and lines[i].startswith('|'):
                i += 1

            # Insert formatted table
            create_tech_stack_table(doc)
            continue

        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                code_content = []
            else:
                in_code_block = False
                # Add code block
                if code_content:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Normal'
                    code_text = '\n'.join(code_content)
                    run = code_para.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    # Add light gray background effect via shading
                    code_para.paragraph_format.left_indent = Inches(0.25)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle headings
        if line.startswith('# ') and not line.startswith('##'):
            heading = doc.add_heading(line[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)

        # Handle horizontal rules
        elif line.strip() == '---':
            para = doc.add_paragraph()
            para.add_run('─' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Handle bold text at line start (like Author:, Tags:, TL;DR:)
        elif line.startswith('**') and ':**' in line:
            para = doc.add_paragraph()
            # Parse bold prefix
            match = re.match(r'\*\*([^*]+)\*\*(.*)$', line)
            if match:
                run = para.add_run(match.group(1))
                run.bold = True
                para.add_run(match.group(2))
            else:
                para.add_run(line)

        # Handle bullet points
        elif line.strip().startswith('- '):
            para = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith('* '):
            para = doc.add_paragraph(line.strip()[2:], style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s(.*)$', line.strip())
            if match:
                para = doc.add_paragraph(match.group(2), style='List Number')

        # Handle regular paragraphs (skip empty lines)
        elif line.strip():
            para = doc.add_paragraph()
            # Process inline formatting
            text = line

            # Handle inline code
            parts = re.split(r'(`[^`]+`)', text)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    # Handle bold
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bold_part in bold_parts:
                        if bold_part.startswith('**') and bold_part.endswith('**'):
                            run = para.add_run(bold_part[2:-2])
                            run.bold = True
                        else:
                            # Handle italic
                            italic_parts = re.split(r'(\*[^*]+\*)', bold_part)
                            for italic_part in italic_parts:
                                if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) > 2:
                                    run = para.add_run(italic_part[1:-1])
                                    run.italic = True
                                else:
                                    # Handle links
                                    link_pattern = r'\[([^\]]+)\]\([^)]+\)'
                                    link_parts = re.split(f'({link_pattern})', italic_part)
                                    for link_part in link_parts:
                                        link_match = re.match(link_pattern, link_part)
                                        if link_match:
                                            run = para.add_run(link_match.group(1))
                                            run.font.color.rgb = RGBColor(59, 130, 246)
                                            run.underline = True
                                        elif link_part:
                                            para.add_run(link_part)

        i += 1

    # Add references section at the end
    add_references_section(doc)

    doc.save(docx_path)
    print(f"Document saved to: {docx_path}")

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(script_dir)

    md_path = os.path.join(root_dir, "DZONE_ARTICLE.md")
    docx_path = os.path.join(root_dir, "DZONE_ARTICLE.docx")

    process_markdown_to_docx(md_path, docx_path)
